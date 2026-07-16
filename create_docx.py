from docx import Document

doc = Document()

doc.add_heading('Society Compliance Copilot - Setup & Run Instructions', 0)

doc.add_heading('1. Start the Backend API', level=1)
doc.add_paragraph('Open a terminal, navigate to the project directory, and run the following commands to activate the virtual environment and start the FastAPI server:')
doc.add_paragraph('cd /home/stark/society-compliance-chatbot', style='Intense Quote')
doc.add_paragraph('source .venv/bin/activate', style='Intense Quote')
doc.add_paragraph('python3 backend/api.py', style='Intense Quote')
doc.add_paragraph('Wait until you see "Uvicorn running on http://0.0.0.0:8000" in the console.')

doc.add_heading('2. Start the Frontend Application', level=1)
doc.add_paragraph('Open a second, new terminal, navigate to the frontend directory, set up the Node.js path (if required), and start the Vite development server:')
doc.add_paragraph('cd /home/stark/society-compliance-chatbot/frontend-react', style='Intense Quote')
doc.add_paragraph('export PATH=/home/stark/society-compliance-chatbot/node-v22.14.0-linux-x64/bin:$PATH', style='Intense Quote')
doc.add_paragraph('npm run dev', style='Intense Quote')

doc.add_heading('3. Access the Application', level=1)
doc.add_paragraph('Once both servers are running, open your web browser and go to:')
doc.add_paragraph('http://localhost:5173/', style='Intense Quote')

doc.save('run_commands.docx')
print("run_commands.docx created successfully!")
